"""
知识库 API
"""

import io
import re
from pathlib import Path

import pandas as pd
from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from pypdf import PdfReader
from sqlalchemy.orm import Session

from backend.database import get_db
from backend.models.models import KnowledgeDocument, KnowledgeChunk
from backend.schemas.schemas import (
    KnowledgeSearchRequest, KnowledgeSearchResponse, KnowledgeDoc,
    KnowledgeUploadRequest,
)
from backend.rag.rag import simple_vector_search

router = APIRouter(prefix="/api/knowledge", tags=["知识库"])
SUPPORTED_FILE_SUFFIXES = {".pdf", ".txt", ".docx", ".md"}


@router.get("")
def list_knowledge(db: Session = Depends(get_db)):
    """List knowledge documents with chunk counts for the management UI."""
    rows = (
        db.query(KnowledgeDocument)
        .order_by(KnowledgeDocument.created_at.desc())
        .all()
    )
    docs = []
    for doc in rows:
        chunks_count = (
            db.query(KnowledgeChunk)
            .filter(KnowledgeChunk.document_id == doc.id)
            .count()
        )
        docs.append({
            "id": doc.id,
            "title": doc.title,
            "doc_type": doc.doc_type,
            "content": doc.content[:220],
            "chunks_count": chunks_count,
            "created_at": doc.created_at.isoformat() if doc.created_at else "",
        })
    return {"docs": docs}


@router.post("/search", response_model=KnowledgeSearchResponse)
def search_knowledge(req: KnowledgeSearchRequest):
    """检索知识库"""
    docs = simple_vector_search(req.query, top_k=req.top_k)
    return KnowledgeSearchResponse(
        docs=[
            KnowledgeDoc(
                id=d.id,
                title=getattr(d, 'title', ''),
                content=d.chunk_text[:500],
                doc_type=getattr(d, 'doc_type', 'general'),
                score=getattr(d, 'score', 0.0),
            )
            for d in docs
        ]
    )


@router.post("/upload")
def upload_knowledge(req: KnowledgeUploadRequest, db: Session = Depends(get_db)):
    """上传知识文档（含自动分块）"""
    return _create_knowledge_document(
        db=db,
        title=req.title,
        doc_type=req.doc_type,
        content=clean_knowledge_text_with_pandas(req.content),
        metadata=req.metadata,
    )


@router.post("/upload-file")
async def upload_knowledge_file(
    file: UploadFile = File(...),
    doc_type: str = Form("general"),
    title: str | None = Form(None),
    db: Session = Depends(get_db),
):
    """上传 PDF / TXT / Word / MD 文件，抽取文本后清洗、切块并入库。"""
    filename = file.filename or "knowledge.txt"
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_FILE_SUFFIXES:
        raise HTTPException(
            status_code=400,
            detail="仅支持 PDF、TXT、DOCX、MD 格式的知识文档",
        )

    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="上传文件为空")

    extracted_text = extract_text_from_upload(filename, raw)
    cleaned_text = clean_knowledge_text_with_pandas(extracted_text)
    if not cleaned_text:
        raise HTTPException(status_code=400, detail="未能从文件中解析出有效文本")

    metadata = {
        "source": "车型库知识文档上传",
        "source_filename": filename,
        "file_suffix": suffix,
        "cleaning": summarize_cleaning(extracted_text, cleaned_text),
    }
    return _create_knowledge_document(
        db=db,
        title=(title or Path(filename).stem).strip(),
        doc_type=doc_type,
        content=cleaned_text,
        metadata=metadata,
    )


def _create_knowledge_document(
    db: Session,
    title: str,
    doc_type: str,
    content: str,
    metadata: dict | None = None,
):
    """保存知识文档并生成知识切片。"""
    metadata = metadata or {}
    doc = KnowledgeDocument(
        title=title,
        doc_type=doc_type,
        content=content,
        doc_metadata=metadata,
    )
    db.add(doc)
    db.flush()

    # 分块（按段落和句子切分）
    chunks = _chunk_text(content)
    for i, chunk_text in enumerate(chunks):
        chunk = KnowledgeChunk(
            document_id=doc.id,
            chunk_text=chunk_text,
            chunk_metadata={"chunk_index": i, **metadata},
        )
        db.add(chunk)

    db.commit()
    return {
        "document_id": doc.id,
        "chunks": len(chunks),
        "message": "文档已上传并分块",
    }


def extract_text_from_upload(filename: str, raw: bytes) -> str:
    """按文件类型抽取原始文本。"""
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md"}:
        return _decode_text(raw)
    if suffix == ".docx":
        document = Document(io.BytesIO(raw))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    raise HTTPException(status_code=400, detail="不支持的文件格式")


def clean_knowledge_text_with_pandas(text: str) -> str:
    """使用 Pandas 做段落级文本清洗，便于后续知识切块和检索。"""
    raw_lines = re.split(r"[\r\n]+", text or "")
    df = pd.DataFrame({"line": raw_lines})
    df["line"] = (
        df["line"]
        .astype(str)
        .str.replace("\u3000", " ", regex=False)
        .str.replace(r"\s+", " ", regex=True)
        .str.strip()
    )
    df = df[df["line"].ne("")]
    df = df.drop_duplicates(subset=["line"], keep="first")
    df["length"] = df["line"].str.len()
    df = df[df["length"] >= 2]
    return "\n".join(df["line"].tolist())


def summarize_cleaning(raw_text: str, cleaned_text: str) -> dict:
    """返回清洗摘要，方便前端和面试讲解。"""
    raw_lines = [line for line in re.split(r"[\r\n]+", raw_text or "") if line.strip()]
    cleaned_lines = [line for line in cleaned_text.splitlines() if line.strip()]
    return {
        "raw_lines": len(raw_lines),
        "cleaned_lines": len(cleaned_lines),
        "removed_lines": max(len(raw_lines) - len(cleaned_lines), 0),
    }


def _decode_text(raw: bytes) -> str:
    """兼容常见中文文档编码。"""
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def _chunk_text(text: str, max_chars: int = 300) -> list[str]:
    """文本分块"""
    chunks = []
    paragraphs = text.split("\n")
    current = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) < max_chars:
            current += para + "\n"
        else:
            if current:
                chunks.append(current.strip())
            # 如果段落本身就很长，按句子拆分
            if len(para) > max_chars:
                import re
                sentences = re.split(r"(?<=[。！？])", para)
                for sent in sentences:
                    if sent:
                        chunks.append(sent.strip())
            else:
                current = para + "\n"

    if current:
        chunks.append(current.strip())

    # 确保没有空块
    return [c for c in chunks if c]
